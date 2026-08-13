from __future__ import annotations

import io
import os
import tempfile
from collections import Counter
from typing import Any

from docx import Document
from faker import Faker

from detectors import detect_pii


# ============================================================
# CONFIGURATION
# ============================================================

fake = Faker("en_IN")


PII_TYPES = [
    "PERSON",
    "EMAIL",
    "PHONE",
    "COMPANY",
    "ADDRESS",
    "SSN",
    "CREDIT_CARD",
    "DOB",
    "IP_ADDRESS",
]


# ============================================================
# FAKE VALUE GENERATOR
# ============================================================

class ReplacementGenerator:

    def __init__(self):
        self.maps = {
            label: {}
            for label in PII_TYPES
        }

    def generate(
        self,
        label: str,
        original: str
    ) -> str:

        if original in self.maps[label]:
            return self.maps[label][original]

        if label == "PERSON":

            value = fake.name()

        elif label == "EMAIL":

            value = fake.email()

        elif label == "PHONE":

            value = (
                "+91 "
                + str(
                    fake.random_number(
                        digits=10,
                        fix_len=True
                    )
                )
            )

        elif label == "COMPANY":

            value = fake.company()

        elif label == "ADDRESS":

            value = fake.address().replace(
                "\n",
                ", "
            )

        elif label == "SSN":

            value = fake.ssn()

        elif label == "CREDIT_CARD":

            value = fake.credit_card_number()

        elif label == "DOB":

            value = fake.date_of_birth(
                minimum_age=25,
                maximum_age=70
            ).strftime("%d/%m/%Y")

        elif label == "IP_ADDRESS":

            value = fake.ipv4()

        else:

            value = "[REDACTED]"

        self.maps[label][original] = value

        return value


# ============================================================
# REPLACE TEXT
# ============================================================

def replace_paragraph(
    paragraph,
    generator: ReplacementGenerator,
    records: list[dict[str, Any]],
    stats: Counter
) -> None:

    original_text = paragraph.text

    if not original_text.strip():
        return

    entities = detect_pii(
        original_text
    )

    if not entities:
        return

    new_text = original_text

    replacements = []

    # IMPORTANT:
    # Replace from right to left because entity offsets
    # refer to the original text.
    for entity in sorted(
        entities,
        key=lambda item: item.start,
        reverse=True
    ):

        replacement = generator.generate(
            entity.label,
            entity.text
        )

        new_text = (
            new_text[:entity.start]
            + replacement
            + new_text[entity.end:]
        )

        replacements.append({
            "label": entity.label,
            "original": entity.text,
            "replacement": replacement,
            "start": entity.start,
            "end": entity.end,
        })

    if new_text != original_text:

        paragraph.text = new_text

    for item in replacements:

        stats[item["label"]] += 1

        records.append(item)


# ============================================================
# PROCESS TABLE
# ============================================================

def process_table(
    table,
    generator,
    records,
    stats
):

    for row in table.rows:

        for cell in row.cells:

            for paragraph in cell.paragraphs:

                replace_paragraph(
                    paragraph,
                    generator,
                    records,
                    stats
                )

            for nested_table in cell.tables:

                process_table(
                    nested_table,
                    generator,
                    records,
                    stats
                )


# ============================================================
# PROCESS DOCUMENT
# ============================================================

def process_document(
    document,
    generator,
    records,
    stats,
    progress_callback=None
):

    total_steps = (
        len(document.paragraphs)
        + len(document.tables)
        + len(document.sections)
    )

    total_steps = max(
        total_steps,
        1
    )

    completed = 0

    # --------------------------------------------------------
    # Paragraphs
    # --------------------------------------------------------

    for paragraph in document.paragraphs:

        replace_paragraph(
            paragraph,
            generator,
            records,
            stats
        )

        completed += 1

        if progress_callback:

            progress_callback(
                min(
                    completed / total_steps,
                    1.0
                )
            )

    # --------------------------------------------------------
    # Tables
    # --------------------------------------------------------

    for table in document.tables:

        process_table(
            table,
            generator,
            records,
            stats
        )

        completed += 1

        if progress_callback:

            progress_callback(
                min(
                    completed / total_steps,
                    1.0
                )
            )

    # --------------------------------------------------------
    # Headers / Footers
    # --------------------------------------------------------

    for section in document.sections:

        for paragraph in section.header.paragraphs:

            replace_paragraph(
                paragraph,
                generator,
                records,
                stats
            )

        for table in section.header.tables:

            process_table(
                table,
                generator,
                records,
                stats
            )

        for paragraph in section.footer.paragraphs:

            replace_paragraph(
                paragraph,
                generator,
                records,
                stats
            )

        for table in section.footer.tables:

            process_table(
                table,
                generator,
                records,
                stats
            )

        completed += 1

        if progress_callback:

            progress_callback(
                min(
                    completed / total_steps,
                    1.0
                )
            )


# ============================================================
# POST-REDACTION VALIDATION
# ============================================================

def validate_redaction(
    document
) -> dict[str, Any]:

    remaining = []

    # --------------------------------------------------------
    # Paragraphs
    # --------------------------------------------------------

    for paragraph in document.paragraphs:

        entities = detect_pii(
            paragraph.text
        )

        for entity in entities:

            remaining.append({
                "label": entity.label,
                "text": entity.text,
                "location": "paragraph"
            })

    # --------------------------------------------------------
    # Tables
    # --------------------------------------------------------

    def inspect_table(table):

        for row in table.rows:

            for cell in row.cells:

                for paragraph in cell.paragraphs:

                    entities = detect_pii(
                        paragraph.text
                    )

                    for entity in entities:

                        remaining.append({
                            "label": entity.label,
                            "text": entity.text,
                            "location": "table"
                        })

                for nested in cell.tables:

                    inspect_table(nested)

    for table in document.tables:

        inspect_table(table)

    # --------------------------------------------------------
    # Headers / Footers
    # --------------------------------------------------------

    for section in document.sections:

        for paragraph in section.header.paragraphs:

            entities = detect_pii(
                paragraph.text
            )

            for entity in entities:

                remaining.append({
                    "label": entity.label,
                    "text": entity.text,
                    "location": "header"
                })

        for paragraph in section.footer.paragraphs:

            entities = detect_pii(
                paragraph.text
            )

            for entity in entities:

                remaining.append({
                    "label": entity.label,
                    "text": entity.text,
                    "location": "footer"
                })

    return {
        "remaining_count": len(remaining),
        "remaining_entities": remaining,
        "passed": len(remaining) == 0
    }


# ============================================================
# MAIN ENGINE
# ============================================================

def redact_docx_bytes(
    file_bytes: bytes,
    progress_callback=None
) -> dict[str, Any]:

    input_stream = io.BytesIO(
        file_bytes
    )

    document = Document(
        input_stream
    )

    paragraph_count = len(
        document.paragraphs
    )

    table_count = len(
        document.tables
    )

    section_count = len(
        document.sections
    )

    generator = ReplacementGenerator()

    records = []

    stats = Counter()

    # --------------------------------------------------------
    # Process
    # --------------------------------------------------------

    process_document(
        document,
        generator,
        records,
        stats,
        progress_callback
    )

    # --------------------------------------------------------
    # Save redacted DOCX into memory
    # --------------------------------------------------------

    output_stream = io.BytesIO()

    document.save(
        output_stream
    )

    output_stream.seek(0)

    output_bytes = (
        output_stream.getvalue()
    )

    # --------------------------------------------------------
    # Validation
    # --------------------------------------------------------

    validation = validate_redaction(
        document
    )

    # --------------------------------------------------------
    # Unique mappings
    # --------------------------------------------------------

    unique_records = []

    seen = set()

    for record in records:

        key = (
            record["label"],
            record["original"],
            record["replacement"]
        )

        if key not in seen:

            seen.add(key)

            unique_records.append(
                record
            )

    # --------------------------------------------------------
    # Document statistics
    # --------------------------------------------------------

    stats_dict = {
        label: int(
            stats.get(label, 0)
        )
        for label in PII_TYPES
    }

    return {

        "document": output_bytes,

        "stats": stats_dict,

        "total_replacements": len(
            records
        ),

        "unique_replacements": len(
            unique_records
        ),

        "records": unique_records,

        "replacement_maps": generator.maps,

        "validation": validation,

        "document_info": {
            "paragraphs": paragraph_count,
            "tables": table_count,
            "sections": section_count,
            "input_size": len(file_bytes),
            "output_size": len(output_bytes),
        }
    }