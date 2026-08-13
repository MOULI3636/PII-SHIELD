import json
import os
from collections import Counter

from docx import Document
from faker import Faker

from detectors import detect_pii


# ============================================================
# CONFIGURATION
# ============================================================

INPUT_FILE = "Red Herring Prospectus.docx"

OUTPUT_DIR = "output"

OUTPUT_FILE = os.path.join(
    OUTPUT_DIR,
    "redacted_prospectus.docx"
)

REPORT_FILE = os.path.join(
    OUTPUT_DIR,
    "detection_report.json"
)


# ============================================================
# FAKER
# ============================================================

fake = Faker("en_IN")


replacement_maps = {
    "PERSON": {},
    "EMAIL": {},
    "PHONE": {},
    "COMPANY": {},
    "ADDRESS": {},
    "SSN": {},
    "CREDIT_CARD": {},
    "DOB": {},
    "IP_ADDRESS": {}
}


# ============================================================
# GENERATE FAKE VALUES
# ============================================================

def generate_fake_value(label, original):

    if original in replacement_maps[label]:
        return replacement_maps[label][original]

    if label == "PERSON":
        value = fake.name()

    elif label == "EMAIL":
        value = fake.email()

    elif label == "PHONE":
        value = "+91 " + str(
            fake.random_number(
                digits=10,
                fix_len=True
            )
        )

    elif label == "COMPANY":
        value = fake.company()

    elif label == "ADDRESS":
        value = fake.address().replace(
            "\n",
            ", "
        )

    elif label == "SSN":
        value = fake.ssn()

    elif label == "CREDIT_CARD":
        value = fake.credit_card_number()

    elif label == "DOB":
        value = fake.date_of_birth(
            minimum_age=25,
            maximum_age=70
        ).strftime("%d/%m/%Y")

    elif label == "IP_ADDRESS":
        value = fake.ipv4()

    else:
        value = "[REDACTED]"

    replacement_maps[label][original] = value

    return value


# ============================================================
# REPLACE ENTITIES
# ============================================================

def replace_text(paragraph, entities):

    if not entities:
        return []

    original_text = paragraph.text

    new_text = original_text

    replacements = []

    # Replace from right to left so positions remain correct.
    for entity in sorted(
        entities,
        key=lambda x: x.start,
        reverse=True
    ):

        fake_value = generate_fake_value(
            entity.label,
            entity.text
        )

        new_text = (
            new_text[:entity.start]
            + fake_value
            + new_text[entity.end:]
        )

        replacements.append({
            "label": entity.label,
            "original": entity.text,
            "replacement": fake_value
        })

    if new_text != original_text:
        paragraph.text = new_text

    return replacements


# ============================================================
# PROCESS PARAGRAPH
# ============================================================

def process_paragraph(
    paragraph,
    stats,
    records
):

    text = paragraph.text

    if not text.strip():
        return

    entities = detect_pii(text)

    if not entities:
        return

    replacements = replace_text(
        paragraph,
        entities
    )

    for item in replacements:

        stats[item["label"]] += 1

        records.append(item)


# ============================================================
# PROCESS TABLE
# ============================================================

def process_table(
    table,
    stats,
    records
):

    for row in table.rows:

        for cell in row.cells:

            for paragraph in cell.paragraphs:

                process_paragraph(
                    paragraph,
                    stats,
                    records
                )

            # Process nested tables
            for nested_table in cell.tables:

                process_table(
                    nested_table,
                    stats,
                    records
                )


# ============================================================
# MAIN PROCESS
# ============================================================

def main():

    print("=" * 70)
    print("PII REDACTION TOOL")
    print("=" * 70)

    print()

    # --------------------------------------------------------
    # Check input file
    # --------------------------------------------------------

    if not os.path.exists(INPUT_FILE):

        print(
            f"ERROR: Cannot find '{INPUT_FILE}'"
        )

        print("\nFiles in current directory:")

        for file in os.listdir("."):
            print(" -", file)

        return

    print(
        f"Input file found: {INPUT_FILE}"
    )

    # --------------------------------------------------------
    # Create output directory
    # --------------------------------------------------------

    os.makedirs(
        OUTPUT_DIR,
        exist_ok=True
    )

    # --------------------------------------------------------
    # Load document
    # --------------------------------------------------------

    print("\nLoading DOCX...")

    document = Document(
        INPUT_FILE
    )

    print(
        f"Paragraphs found: "
        f"{len(document.paragraphs)}"
    )

    print(
        f"Tables found: "
        f"{len(document.tables)}"
    )

    stats = Counter()

    records = []

    # --------------------------------------------------------
    # Process normal paragraphs
    # --------------------------------------------------------

    print("\nProcessing paragraphs...")

    for paragraph in document.paragraphs:

        process_paragraph(
            paragraph,
            stats,
            records
        )

    # --------------------------------------------------------
    # Process tables
    # --------------------------------------------------------

    print("Processing tables...")

    for table in document.tables:

        process_table(
            table,
            stats,
            records
        )

    # --------------------------------------------------------
    # Process headers and footers
    # --------------------------------------------------------

    print("Processing headers and footers...")

    for section in document.sections:

        # Header paragraphs
        for paragraph in section.header.paragraphs:

            process_paragraph(
                paragraph,
                stats,
                records
            )

        # Header tables
        for table in section.header.tables:

            process_table(
                table,
                stats,
                records
            )

        # Footer paragraphs
        for paragraph in section.footer.paragraphs:

            process_paragraph(
                paragraph,
                stats,
                records
            )

        # Footer tables
        for table in section.footer.tables:

            process_table(
                table,
                stats,
                records
            )

    # --------------------------------------------------------
    # Save redacted document
    # --------------------------------------------------------

    print("\nSaving redacted document...")

    document.save(
        OUTPUT_FILE
    )

    # --------------------------------------------------------
    # Remove duplicate records
    # --------------------------------------------------------

    unique_records = []

    seen = set()

    for record in records:

        key = (
            record["label"],
            record["original"],
            record["replacement"]
        )

        if key not in seen:

            seen.add(key)

            unique_records.append(
                record
            )

    # --------------------------------------------------------
    # Detection report
    # --------------------------------------------------------

    report = {

        "input_file": INPUT_FILE,

        "output_file": OUTPUT_FILE,

        "total_replacements": len(records),

        "unique_replacements": len(
            unique_records
        ),

        "counts_by_type": dict(
            stats
        ),

        "replacement_mapping": unique_records
    }

    print(
        "Writing detection report..."
    )

    with open(
        REPORT_FILE,
        "w",
        encoding="utf-8"
    ) as file:

        json.dump(
            report,
            file,
            indent=4,
            ensure_ascii=False
        )

    # --------------------------------------------------------
    # Final output
    # --------------------------------------------------------

    print()

    print("=" * 70)
    print("REDACTION COMPLETE")
    print("=" * 70)

    print("\nPII detected:")

    if stats:

        for label in sorted(stats):

            print(
                f"  {label:<15}: "
                f"{stats[label]}"
            )

    else:

        print(
            "  WARNING: No PII detected."
        )

    print()

    print(
        f"Total replacements : "
        f"{len(records)}"
    )

    print(
        f"Unique replacements: "
        f"{len(unique_records)}"
    )

    print()

    print("Redacted document:")
    print(f"  {OUTPUT_FILE}")

    print()

    print("Detection report:")
    print(f"  {REPORT_FILE}")

    print()

    print("SUCCESS")


# ============================================================
# ENTRY POINT
# ============================================================

if __name__ == "__main__":
    main()